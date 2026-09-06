import http.server
import socketserver
import os
import json
import io
import re
import time
import urllib.request
import urllib.error
from urllib.parse import unquote

PORT = 3001
MAX_UPLOAD = 50 * 1024 * 1024      # 50 MB body cap
MAX_EXTRACT_CHARS = 100_000        # extraction text cap (protects LLM context)

EXTRACTABLE = {".pdf", ".xlsx", ".xls", ".docx", ".txt", ".csv", ".json", ".md"}

# ── Voice read-aloud (2026-08-31): Qwen3-TTS cloned voices, proxied from the
#    local qwen3-tts microservice (port 5093) so the browser never needs CORS.
#    ~/ forms only — this file ships to the public R-SITES repo (no system paths).
Q3_BASE = "http://127.0.0.1:5093"
VOICES_DIR = os.environ.get("VOICES_DIR", os.path.expanduser("~/kokoro-playground/cloned_voices"))
VOICES_JSON = os.path.join(VOICES_DIR, "voices.json")
TTS_TIMEOUT = 120          # per-chunk synth cap (client chunks long replies)
TTS_MAX_TEXT = 6000        # per-request text cap (sanity guard)

# ── Agent proxy (2026-09-04): the browser chats with connected AGENTS (Hermes
#    gateway on :8642, Pi / DS Harness / custom OpenAI-compatible servers) by
#    POSTing to THIS server — the gateway has no CORS, and agent keys must never
#    ship in the public repo. The Hermes key lives ONLY in the gitignored
#    .agent-keys.json next to this file (or HERMES_API_KEY env). Keys the user
#    types for other agents ride the request body over localhost (runtime data,
#    never baked into any shipped file).
AGENT_PING_TIMEOUT = 4     # reachability probe cap
AGENT_CHAT_TIMEOUT = None  # no cap — agent tasks can run 30+ minutes
_MODEL_CACHE = {}          # base url -> (model_id, fetched_ts) — default model resolution
# Runs protocol (2026-09-05): agent mode chats via POST /v1/runs — the gateway
# returns a run_id immediately, then GET /v1/runs/{id}/events streams typed
# lifecycle events (message.delta, tool.started/completed, reasoning.available,
# approval.request, run.completed/failed/cancelled). Approvals resolve via
# POST /v1/runs/{id}/approval {choice: once|session|always|deny}; steering via
# /steer {input}; stopping via /stop. This is the surface the Hermes CLI-like
# frontends use — chat/completions never emits approvals (verified in source).

# ── Media proxy (2026-09-05): inline agent-generated images in chat responses.
#    The Hermes gateway references local files in the streamed text (MEDIA:/abs/path,
#    ~/…/x.png, ComfyUI/output/x.png) but a browser cannot read the local disk, so the
#    client renders <img src="/api/media?path=…"> and we resolve + serve the bytes here.
#    Resolution mirrors R1 Connect's resolve_image_path() exactly (battle-tested against
#    the same gateway). Relative refs resolve against the R1 base roots in R1's order.
MEDIA_EXTS = {".png", ".jpg", ".jpeg", ".webp", ".gif", ".bmp",
              ".mp4", ".mov", ".m4v", ".webm", ".mp3", ".wav", ".m4a",
              ".aac", ".ogg", ".flac", ".opus"}
MEDIA_IMG_EXTS = {".png", ".jpg", ".jpeg", ".webp", ".gif", ".bmp"}
MEDIA_CT = {".png": "image/png", ".jpg": "image/jpeg", ".jpeg": "image/jpeg",
            ".webp": "image/webp", ".gif": "image/gif", ".bmp": "image/bmp",
            ".mp4": "video/mp4", ".mov": "video/quicktime", ".m4v": "video/x-m4v",
            ".webm": "video/webm", ".mp3": "audio/mpeg", ".wav": "audio/wav",
            ".m4a": "audio/mp4", ".aac": "audio/aac", ".ogg": "audio/ogg",
            ".flac": "audio/flac", ".opus": "audio/ogg"}
MAX_MEDIA_BYTES = 8 * 1024 * 1024   # 8 MB cap per image (Krea/ComfyUI PNGs run 1-3 MB)
MAX_MEDIA_VIDEO_BYTES = 256 * 1024 * 1024  # 256 MB cap per video/audio file
MEDIA_BASE_PATHS = [                # relative refs resolve against these, in order
    os.path.expanduser("~/ComfyUI/output"),
    os.path.expanduser("~/ComfyUI/output/img2ltx"),
    os.path.expanduser("~"),
    os.getcwd(),
]


def resolve_media_path(raw):
    """Resolve a media token to a real local file, or None.

    Mirrors r1-ws-server.resolve_image_path(): strip a MEDIA: prefix, try the
    absolute path, then the token (minus any leading .//) against each base root.
    Strips trailing prose punctuation so 'x.png.' and 'x.png)' still resolve.
    """
    if not raw:
        return None
    path = raw.strip().strip(".,;:!?)]}>\"'")
    if path.lower().startswith(("http://", "https://", "data:", "file:", "blob:", "ftp://")):
        return None
    if path[:6].upper() == "MEDIA:":
        path = path[6:]
    path = os.path.expanduser(path)
    if not os.path.isabs(path):
        clean = re.sub(r"^\.?[/\\]", "", path)
        for base in MEDIA_BASE_PATHS:
            cand = os.path.join(base, clean)
            if os.path.isfile(cand):
                return cand
        return None
    return path if os.path.isfile(path) else None


def _load_agent_keys():
    try:
        with open(os.path.join(os.path.dirname(os.path.abspath(__file__)),
                               ".agent-keys.json"), "r", encoding="utf-8") as f:
            return json.load(f) or {}
    except Exception:
        return {}


def agent_secret(agent_id, req_key):
    """Browser-supplied key wins; Hermes falls back to the local key file/env."""
    if req_key:
        return req_key
    if agent_id == "hermes":
        keys = _load_agent_keys()
        return keys.get("hermes") or os.environ.get("HERMES_API_KEY", "")
    return ""


def agent_urls(base):
    """Split a base URL into chat + models endpoints (OpenAI-compatible layout)."""
    base = (base or "").strip().rstrip("/")
    if base.endswith("/chat/completions"):
        stem = base[: -len("/chat/completions")]
        return base, stem + "/models"
    if base.endswith("/v1"):
        return base + "/chat/completions", base + "/models"
    return base + "/v1/chat/completions", base + "/v1/models"


def agent_api_base(base):
    """Normalize a base URL to its API root (no /v1, no /chat/completions)."""
    base = (base or "").strip().rstrip("/")
    if base.endswith("/chat/completions"):
        base = base[: -len("/chat/completions")]
    if base.endswith("/v1"):
        base = base[: -len("/v1")]
    return base.rstrip("/")


def agent_runs_url(base):
    """Runs-protocol endpoints under an agent base URL."""
    root = agent_api_base(base)
    return root + "/v1/runs"


def agent_open(url, key, method="GET", body=None, headers=None, timeout=None):
    """timeout=None = no cap (long agent tasks); ping passes a small cap."""
    hdrs = dict(headers or {})
    if key:
        hdrs["Authorization"] = "Bearer " + key
    data = None
    if body is not None:
        data = json.dumps(body).encode()
        hdrs["Content-Type"] = "application/json"
    req = urllib.request.Request(url, data=data, headers=hdrs, method=method)
    return urllib.request.urlopen(req, timeout=timeout)


def extract_text(filename, data):
    ext = os.path.splitext(filename)[1].lower()
    text = ""
    note = ""
    if ext == ".pdf":
        from pypdf import PdfReader
        r = PdfReader(io.BytesIO(data))
        pages = [p.extract_text() or "" for p in r.pages]
        text = "\n\n".join(f"[page {i + 1}]\n{t}" for i, t in enumerate(pages))
        note = f"{len(r.pages)} pages"
    elif ext == ".xlsx":
        from openpyxl import load_workbook
        wb = load_workbook(io.BytesIO(data), read_only=True, data_only=True)
        parts = []
        for ws in wb.worksheets:
            rows = []
            for row in ws.iter_rows(values_only=True):
                rows.append("\t".join("" if v is None else str(v) for v in row))
            parts.append(f"[sheet: {ws.title}]\n" + "\n".join(rows))
        text = "\n\n".join(parts)
        note = f"{len(wb.worksheets)} sheets"
        wb.close()
    elif ext == ".xls":
        import xlrd
        wb = xlrd.open_workbook(file_contents=data)
        parts = []
        for sh in wb.sheets():
            rows = []
            for r in range(sh.nrows):
                rows.append("\t".join(str(sh.cell_value(r, c)) for c in range(sh.ncols)))
            parts.append(f"[sheet: {sh.name}]\n" + "\n".join(rows))
        text = "\n\n".join(parts)
        note = f"{wb.nsheets} sheets"
    elif ext == ".docx":
        from docx import Document
        d = Document(io.BytesIO(data))
        paras = [p.text for p in d.paragraphs]
        tables = []
        for t in d.tables:
            for row in t.rows:
                tables.append("\t".join(c.text for c in row.cells))
        text = "\n".join(paras + tables)
        note = f"{len(d.paragraphs)} paragraphs"
    else:  # txt / csv / json / md — client usually reads these, but support anyway
        text = data.decode("utf-8", errors="replace")
    if len(text) > MAX_EXTRACT_CHARS:
        text = text[:MAX_EXTRACT_CHARS] + "\n…[truncated]"
    return text, note


# ── Voice helpers ────────────────────────────────────────────────────────────

def q3_health():
    try:
        with urllib.request.urlopen(Q3_BASE + "/health", timeout=3) as r:
            return r.status == 200
    except Exception:
        return False


def load_voices():
    try:
        with open(VOICES_JSON, "r", encoding="utf-8") as f:
            data = json.load(f)
    except Exception:
        return {}
    return {n: (m if isinstance(m, dict) else {}) for n, m in (data or {}).items()}


def tts_synthesize(voice_name, text):
    """Clone-speak `text` with the stored reference of `voice_name` (Qwen3-TTS).
    Returns (audio_bytes, None) or (None, (error_dict, http_code))."""
    voices = load_voices()
    meta = voices.get(voice_name)
    if not meta:
        return None, ({"error": f"voice not found: {voice_name}"}, 404)
    filename = meta.get("filename")
    if not filename:
        return None, ({"error": f"voice {voice_name} has no audio file"}, 500)
    ref_path = os.path.join(VOICES_DIR, os.path.basename(filename))
    if not os.path.isfile(ref_path):
        return None, ({"error": f"reference audio missing: {filename}"}, 500)
    with open(ref_path, "rb") as f:
        ref_bytes = f.read()
    boundary = "----hermes" + os.urandom(8).hex()
    parts = []
    parts.append(f"--{boundary}\r\nContent-Disposition: form-data; name=\"audio\"; "
                 f"filename=\"{os.path.basename(filename)}\"\r\n"
                 f"Content-Type: audio/wav\r\n\r\n".encode())
    parts.append(ref_bytes)
    parts.append(b"\r\n")
    parts.append(f"--{boundary}\r\nContent-Disposition: form-data; name=\"text\"\r\n\r\n{text}\r\n".encode())
    ref_text = meta.get("ref_text")
    if ref_text:
        parts.append(f"--{boundary}\r\nContent-Disposition: form-data; name=\"ref_text\"\r\n\r\n{ref_text}\r\n".encode())
    parts.append(f"--{boundary}--\r\n".encode())
    body = b"".join(parts)
    req = urllib.request.Request(Q3_BASE + "/clone", data=body,
                                 headers={"Content-Type": f"multipart/form-data; boundary={boundary}"})
    try:
        with urllib.request.urlopen(req, timeout=TTS_TIMEOUT) as r:
            audio = r.read()
    except urllib.error.HTTPError as e:
        return None, ({"error": f"qwen3 backend error: HTTP {e.code}"}, 502)
    except Exception as e:
        return None, ({"error": f"qwen3 backend unreachable: {e}"}, 502)
    if not audio.startswith(b"RIFF"):
        return None, ({"error": "qwen3 backend returned non-audio",
                       "detail": audio[:300].decode("utf-8", "replace")}, 502)
    return audio, None


class Handler(http.server.SimpleHTTPRequestHandler):
    def __init__(self, *args, **kwargs):
        super().__init__(*args, directory=os.path.dirname(os.path.abspath(__file__)), **kwargs)

    def end_headers(self):
        # no-store: the browser must NEVER keep index.html (all CSS/JS is
        # inline), so every plain reload re-fetches — no Ctrl+Shift+R, and
        # David never has to clear browsing data (which wiped his chats once).
        self.send_header("Cache-Control", "no-store, no-cache, must-revalidate")
        self.send_header("Pragma", "no-cache")
        self.send_header("Expires", "0")
        super().end_headers()

    def log_message(self, format, *args):
        # keep the access log ON (2026-08-27): needed to prove whether a
        # browser actually fetched the page (stale-tab debugging) — the old
        # `pass` made "why is my profile showing an old page?" undiagnosable.
        # Guard the write: if the parent session died and stdout is an
        # orphaned pipe, an uncaught BrokenPipeError here kills the request
        # handler (empty reply on every request — hit live 2026-08-27).
        try:
            print(f"[{self.log_date_time_string()}] {self.client_address[0]} {format % args}", flush=True)
        except (BrokenPipeError, OSError):
            pass

    def do_GET(self):
        if self.path.split("?")[0] == "/api/media":
            self._handle_media()
            return
        if self.path.split("?")[0] == "/api/voices":
            available = q3_health()
            items = [{"name": n,
                      "filename": m.get("filename", ""),
                      "engine": m.get("engine", ""),
                      "created_at": m.get("created_at", "")}
                     for n, m in load_voices().items()]
            self._json({"available": available, "voices": items})
            return
        super().do_GET()

    def _handle_media(self):
        """Serve a resolved local media file (image/video/audio) for the chat.

        The query 'path' is the raw token from the model's text (MEDIA:/abs/path,
        ~/rel.mp4, ComfyUI/output/x.png). Resolution rules live in
        resolve_media_path(); this handler only guards ext/size and streams bytes.
        Supports single-range requests (Accept-Ranges: bytes, 206) so <video> can
        seek without downloading the whole file.
        """
        params = {}
        qs = self.path.split("?", 1)
        if len(qs) > 1:
            for kv in qs[1].split("&"):
                if "=" in kv:
                    k, v = kv.split("=", 1)
                    params[k] = unquote(v)
        resolved = resolve_media_path((params.get("path") or "").strip())
        if not resolved:
            self._json({"error": "media not found"}, 404)
            return
        ext = os.path.splitext(resolved)[1].lower()
        if ext not in MEDIA_EXTS:
            self._json({"error": "not an image"}, 404)
            return
        try:
            size = os.path.getsize(resolved)
        except OSError as e:
            self._json({"error": str(e)}, 404)
            return
        cap = MAX_MEDIA_BYTES if ext in MEDIA_IMG_EXTS else MAX_MEDIA_VIDEO_BYTES
        if size > cap:
            self._json({"error": f"media too large ({size} bytes)"}, 413)
            return
        ctype = MEDIA_CT[ext]
        start, end = 0, size - 1
        rng = self.headers.get("Range")
        if rng:
            m = re.match(r"bytes=(\d*)-(\d*)$", rng.strip())
            if not m:
                self._json({"error": "invalid range"}, 416)
                return
            if m.group(1):
                start = int(m.group(1))
            if m.group(2):
                end = min(int(m.group(2)), size - 1)
            if start > end or start >= size:
                self.send_response(416)
                self.send_header("Content-Range", f"bytes */{size}")
                self.end_headers()
                return
            partial = True
        else:
            partial = False
        length = end - start + 1
        self.send_response(206 if partial else 200)
        self.send_header("Content-Type", ctype)
        self.send_header("Accept-Ranges", "bytes")
        self.send_header("Content-Length", str(length))
        if partial:
            self.send_header("Content-Range", f"bytes {start}-{end}/{size}")
        self.end_headers()
        try:
            with open(resolved, "rb") as f:
                if start:
                    f.seek(start)
                remaining = length
                while remaining > 0:
                    chunk = f.read(min(65536, remaining))
                    if not chunk:
                        break
                    self.wfile.write(chunk)
                    remaining -= len(chunk)
        except (BrokenPipeError, ConnectionResetError):
            pass
        except OSError as e:
            try:
                self._json({"error": str(e)}, 404)
            except Exception:
                pass

    def do_POST(self):
        path = self.path.split("?")[0]
        if path == "/api/tts":
            self._handle_tts()
            return
        if path == "/api/agent/ping":
            self._handle_agent_ping()
            return
        if path == "/api/agent/models":
            self._handle_agent_models()
            return
        if path == "/api/agent/chat":
            self._handle_agent_chat()
            return
        if path == "/api/agent/run":
            self._handle_agent_run()
            return
        if path == "/api/agent/runevents":
            self._handle_agent_runevents()
            return
        if path == "/api/agent/runaction":
            self._handle_agent_runaction()
            return
        if path != "/api/extract":
            self._json({"error": "not found"}, 404)
            return
        length = int(self.headers.get("Content-Length", 0) or 0)
        if length <= 0 or length > MAX_UPLOAD:
            self._json({"error": f"body must be 1..{MAX_UPLOAD} bytes"},
                       413 if length > MAX_UPLOAD else 400)
            return
        filename = self.headers.get("X-Filename", "file")
        data = self.rfile.read(length)
        ext = os.path.splitext(filename)[1].lower()
        if ext not in EXTRACTABLE:
            self._json({"error": f"unsupported type: {ext or '(none)'}"}, 415)
            return
        try:
            text, note = extract_text(filename, data)
            self._json({"text": text, "filename": filename, "note": note,
                        "chars": len(text),
                        "truncated": len(text) >= MAX_EXTRACT_CHARS})
        except Exception as e:
            self._json({"error": f"extract failed: {e}"}, 422)

    def _handle_tts(self):
        length = int(self.headers.get("Content-Length", 0) or 0)
        if length <= 0 or length > MAX_UPLOAD:
            self._json({"error": "empty or too large body"},
                       413 if length > MAX_UPLOAD else 400)
            return
        try:
            req = json.loads(self.rfile.read(length).decode("utf-8", "replace"))
        except Exception:
            self._json({"error": "invalid JSON"}, 400)
            return
        voice = (req.get("voice") or "").strip()
        text = (req.get("text") or "").strip()
        if not voice or not text:
            self._json({"error": "voice and text are required"}, 400)
            return
        if len(text) > TTS_MAX_TEXT:
            self._json({"error": f"text too long (max {TTS_MAX_TEXT} chars)"}, 413)
            return
        audio, err = tts_synthesize(voice, text)
        if err:
            self._json(err[0], err[1])
            return
        assert audio is not None  # guaranteed: err is None on this path
        self.send_response(200)
        self.send_header("Content-Type", "audio/wav")
        self.send_header("Content-Length", str(len(audio)))
        self.end_headers()
        self.wfile.write(audio)

    def _read_json_body(self):
        length = int(self.headers.get("Content-Length", 0) or 0)
        if length <= 0 or length > MAX_UPLOAD:
            return None, ({"error": "empty or too large body"},
                          413 if length > MAX_UPLOAD else 400)
        try:
            return json.loads(self.rfile.read(length).decode("utf-8", "replace")), None
        except Exception:
            return None, ({"error": "invalid JSON"}, 400)

    def _handle_agent_ping(self):
        """Probe an agent's /v1/models so the client can show connected/offline.
        Always answers 200 with {ok:bool} — the CLIENT reads the flag, so a dead
        agent is a JSON verdict, not a fetch failure."""
        req, err = self._read_json_body()
        if err:
            self._json(err[0], err[1])
            return
        assert req is not None
        agent_id = (req.get("id") or "").strip()
        base = (req.get("url") or "").strip()
        if not base:
            self._json({"ok": False, "error": "no server URL configured"})
            return
        _, models_url = agent_urls(base)
        key = agent_secret(agent_id, (req.get("key") or "").strip())
        try:
            with agent_open(models_url, key, timeout=AGENT_PING_TIMEOUT) as r:
                head = r.read(400).decode("utf-8", "replace")
            self._json({"ok": True, "status": r.status, "label": head[:300]})
        except urllib.error.HTTPError as e:
            detail = ""
            try:
                detail = e.read(300).decode("utf-8", "replace")
            except Exception:
                pass
            self._json({"ok": False, "status": e.code, "error": detail or f"HTTP {e.code}"})
        except Exception as e:
            self._json({"ok": False, "error": str(e)[:200]})

    def _handle_agent_models(self):
        """List an OpenAI-compatible endpoint's models (DeepSeek / OpenRouter /
        custom) so the client's cloud group in the header pill can render model
        rows. Returns {ok, models:[{id, ctx}]} — ctx from the provider's own
        context fields when present (token-bar limit)."""
        req, err = self._read_json_body()
        if err:
            self._json(err[0], err[1])
            return
        assert req is not None
        agent_id = (req.get("id") or "").strip()
        base = (req.get("url") or "").strip()
        if not base:
            self._json({"ok": False, "error": "no server URL configured"})
            return
        _, models_url = agent_urls(base)
        key = agent_secret(agent_id, (req.get("key") or "").strip())
        try:
            with agent_open(models_url, key, timeout=AGENT_PING_TIMEOUT) as r:
                if r.status != 200:
                    detail = r.read(300).decode("utf-8", "replace")
                    self._json({"ok": False, "status": r.status,
                                "error": (detail or f"HTTP {r.status}")[:300]})
                    return
                data = json.loads(r.read(400_000).decode("utf-8", "replace"))
            out = []
            for m in (data.get("data") or [])[:500]:
                mid = (m or {}).get("id")
                if not mid:
                    continue
                meta = (m.get("meta") or {}) if isinstance(m.get("meta"), dict) else {}
                ctx = (m.get("max_model_len") or m.get("context_length")
                       or m.get("max_context_length") or meta.get("n_ctx")
                       or meta.get("n_ctx_train") or 0)
                out.append({"id": mid, "ctx": int(ctx) if ctx else 0})
            self._json({"ok": True, "models": out})
        except urllib.error.HTTPError as e:
            detail = ""
            try:
                detail = e.read(300).decode("utf-8", "replace")
            except Exception:
                pass
            self._json({"ok": False, "status": e.code,
                        "error": (detail or f"HTTP {e.code}")[:300]})
        except Exception as e:
            self._json({"ok": False, "error": str(e)[:200]})

    def _handle_agent_chat(self):
        """Stream an OpenAI-compatible agent chat (SSE) back to the browser.
        The body carries the agent target + session id; keys resolve server-side.
        Chunks are forwarded verbatim (event: / data: lines intact) so the client
        sees the EXACT wire stream — no re-parsing here."""
        req, err = self._read_json_body()
        if err:
            self._json(err[0], err[1])
            return
        assert req is not None
        agent_id = (req.get("id") or "").strip()
        base = (req.get("url") or "").strip()
        if not base:
            self._json({"error": "agent has no server URL — set it in Settings → Connect Agent"}, 400)
            return
        chat_url, _ = agent_urls(base)
        key = agent_secret(agent_id, (req.get("key") or "").strip())
        payload = {
            "model": (req.get("model") or "").strip(),
            "messages": req.get("messages") or [],
            "stream": True,
        }
        if not payload["model"]:
            # Plain OpenAI-compatible agents (DeepSeek/Pi/custom) REQUIRE a
            # model. Resolve a default from their /v1/models (cached 120s);
            # DeepSeek falls back to its chat model if the probe fails.
            payload["model"] = self._resolve_agent_model(base, key)
        if not payload["model"]:
            payload.pop("model", None)  # gateway-style: let the server decide
        if req.get("max_tokens"):
            payload["max_tokens"] = int(req["max_tokens"])
        if req.get("temperature") is not None:
            payload["temperature"] = float(req["temperature"])
        if req.get("top_p") is not None:
            payload["top_p"] = float(req["top_p"])
        if req.get("include_usage"):
            # cloud chats (DeepSeek / OpenRouter): exact usage in the final chunk
            payload["stream_options"] = {"include_usage": True}
        headers = {}
        if key:
            headers["Authorization"] = "Bearer " + key
        sess = (req.get("session_id") or "").strip()
        if sess:
            headers["X-Hermes-Session-Id"] = sess
        try:
            upstream = agent_open(chat_url, "", method="POST", body=payload,
                                  headers=headers, timeout=AGENT_CHAT_TIMEOUT)
        except urllib.error.HTTPError as e:
            detail = ""
            try:
                detail = e.read(400).decode("utf-8", "replace")
            except Exception:
                pass
            self._json({"error": detail or f"agent HTTP {e.code}", "status": e.code}, 502)
            return
        except Exception as e:
            self._json({"error": f"agent unreachable: {str(e)[:200]}"}, 502)
            return
        try:
            self.send_response(200)
            self.send_header("Content-Type", "text/event-stream")
            self.send_header("Cache-Control", "no-cache")
            self.send_header("Connection", "close")
            self.end_headers()
            with upstream:
                while True:
                    chunk = upstream.read(8192)
                    if not chunk:
                        break
                    self.wfile.write(chunk)
                    self.wfile.flush()
        except (BrokenPipeError, ConnectionResetError):
            pass  # client walked away mid-stream — normal
        except Exception as e:
            print(f"[agent-chat] stream error: {e}", flush=True)

    def _resolve_agent_model(self, base, key):
        """Pick a default model id from the agent's /v1/models (cached 120s)."""
        now = time.time()
        cached = _MODEL_CACHE.get(base)
        if cached and now - cached[1] < 120:
            return cached[0]
        try:
            _, models_url = agent_urls(base)
            with agent_open(models_url, key, timeout=6) as r:
                data = json.loads(r.read(200_000).decode("utf-8", "replace"))
            mid = (data.get("data") or [{}])[0].get("id")
            if mid:
                _MODEL_CACHE[base] = (mid, now)
                return mid
        except Exception:
            pass
        if "deepseek" in base:
            return "deepseek-chat"
        return None

    def _agent_upstream_err(self, e):
        """Read an upstream HTTPError body for a friendlier client error."""
        detail = ""
        if isinstance(e, urllib.error.HTTPError):
            try:
                detail = e.read(400).decode("utf-8", "replace")
            except Exception:
                pass
            return {"error": detail or f"agent HTTP {e.code}", "status": e.code}
        return {"error": f"agent unreachable: {str(e)[:200]}"}

    def _handle_agent_run(self):
        """POST /api/agent/run — start a /v1/runs agent run; return its run_id.
        Body: {id, url, key, input, session_id?, instructions?, model?}."""
        req, err = self._read_json_body()
        if err:
            self._json(err[0], err[1])
            return
        assert req is not None
        agent_id = (req.get("id") or "").strip()
        base = (req.get("url") or "").strip()
        if not base:
            self._json({"error": "agent has no server URL — set it in Settings → Connect Agent"}, 400)
            return
        key = agent_secret(agent_id, (req.get("key") or "").strip())
        payload = {"input": req.get("input") or [{"role": "user", "content": ""}]}
        sess = (req.get("session_id") or "").strip()
        if sess:
            payload["session_id"] = sess
        if req.get("instructions"):
            payload["instructions"] = req["instructions"]
        if req.get("model"):
            payload["model"] = req["model"]
        try:
            with agent_open(agent_runs_url(base), key, method="POST",
                            body=payload, timeout=AGENT_CHAT_TIMEOUT) as r:
                body = r.read().decode("utf-8", "replace")
            self.send_response(r.status)
            self.send_header("Content-Type", "application/json")
            self.send_header("Content-Length", str(len(body.encode())))
            self.end_headers()
            self.wfile.write(body.encode())
        except urllib.error.HTTPError as e:
            # 404/405 = the agent's server does NOT speak /v1/runs (plain
            # OpenAI-compatible endpoints like DeepSeek) — RELAY the code so the
            # client's designed fallback to the chat-completions turn fires.
            # Any other upstream HTTP error is a real failure -> 502.
            if e.code in (404, 405):
                detail = ""
                try:
                    detail = e.read(400).decode("utf-8", "replace")
                except Exception:
                    pass
                self._json(json.loads(detail) if detail.startswith("{") else {"error": detail},
                           e.code)
                return
            self._json(self._agent_upstream_err(e), 502)
        except Exception as e:
            self._json(self._agent_upstream_err(e), 502)

    def _handle_agent_runevents(self):
        """POST /api/agent/runevents — stream GET /v1/runs/{run_id}/events SSE.
        Body: {id, url, key, run_id}. Events are forwarded verbatim."""
        req, err = self._read_json_body()
        if err:
            self._json(err[0], err[1])
            return
        assert req is not None
        agent_id = (req.get("id") or "").strip()
        base = (req.get("url") or "").strip()
        run_id = (req.get("run_id") or "").strip()
        if not base or not run_id:
            self._json({"error": "url and run_id are required"}, 400)
            return
        key = agent_secret(agent_id, (req.get("key") or "").strip())
        events_url = agent_runs_url(base) + "/" + run_id + "/events"
        try:
            upstream = agent_open(events_url, key, timeout=AGENT_CHAT_TIMEOUT)
        except Exception as e:
            self._json(self._agent_upstream_err(e), 502)
            return
        try:
            self.send_response(200)
            self.send_header("Content-Type", "text/event-stream")
            self.send_header("Cache-Control", "no-cache")
            self.send_header("Connection", "close")
            self.end_headers()
            with upstream:
                while True:
                    chunk = upstream.read(8192)
                    if not chunk:
                        break
                    self.wfile.write(chunk)
                    self.wfile.flush()
        except (BrokenPipeError, ConnectionResetError):
            pass  # client walked away mid-stream — the run keeps going server-side
        except Exception as e:
            print(f"[agent-runevents] stream error: {e}", flush=True)

    def _handle_agent_runaction(self):
        """POST /api/agent/runaction — approval | steer | stop on a live run.
        Body: {id, url, key, run_id, action, choice?, message?}."""
        req, err = self._read_json_body()
        if err:
            self._json(err[0], err[1])
            return
        assert req is not None
        agent_id = (req.get("id") or "").strip()
        base = (req.get("url") or "").strip()
        run_id = (req.get("run_id") or "").strip()
        action = (req.get("action") or "").strip().lower()
        if not base or not run_id or action not in ("approval", "steer", "stop"):
            self._json({"error": "url, run_id and action (approval|steer|stop) are required"}, 400)
            return
        key = agent_secret(agent_id, (req.get("key") or "").strip())
        url = agent_runs_url(base) + "/" + run_id + "/" + action
        body = {}
        if action == "approval":
            body["choice"] = (req.get("choice") or "").strip().lower()
            if body["choice"] not in ("once", "session", "always", "deny"):
                self._json({"error": "approval choice must be once|session|always|deny"}, 400)
                return
        elif action == "steer":
            body["input"] = (req.get("message") or "").strip()
            if not body["input"]:
                self._json({"error": "steer needs a message"}, 400)
                return
        try:
            with agent_open(url, key, method="POST", body=body,
                            timeout=AGENT_CHAT_TIMEOUT) as r:
                out = r.read().decode("utf-8", "replace")
            self.send_response(r.status)
            self.send_header("Content-Type", "application/json")
            self.send_header("Content-Length", str(len(out.encode())))
            self.end_headers()
            self.wfile.write(out.encode())
        except Exception as e:
            self._json(self._agent_upstream_err(e), 502)

    def _json(self, obj, code=200):
        body = json.dumps(obj).encode()
        self.send_response(code)
        self.send_header("Content-Type", "application/json")
        self.send_header("Content-Length", str(len(body)))
        self.end_headers()
        self.wfile.write(body)


with http.server.ThreadingHTTPServer(("", PORT), Handler) as httpd:
    print(f"Chat server running on port {PORT}")
    httpd.serve_forever()
